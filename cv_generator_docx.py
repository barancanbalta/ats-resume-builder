from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO


class ATSResumeDocx:
    """
    Generate ATS-friendly DOCX resume using python-docx
    """

    def __init__(self, language='tr'):
        self.doc = Document()
        self.language = language
        self.setup_document()

        # Localized section headers
        self.labels = {
            'tr': {
                'SUMMARY': 'ÖZET',
                'EXPERIENCE': 'İŞ DENEYİMİ',
                'EDUCATION': 'EĞİTİM',
                'SKILLS': 'TEKNİK BECERİLER'
            },
            'en': {
                'SUMMARY': 'PROFESSIONAL SUMMARY',
                'EXPERIENCE': 'PROFESSIONAL EXPERIENCE',
                'EDUCATION': 'EDUCATION',
                'SKILLS': 'TECHNICAL SKILLS'
            }
        }

    def setup_document(self):
        """Configure document margins, styles and metadata"""
        # Set margins (0.75 inches)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Set default font for all paragraphs
        style = self.doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

    def set_metadata(self, name, summary, skills_dict):
        """Set DOCX Core Properties for ATS scanners"""
        core_props = self.doc.core_properties
        core_props.author = name
        core_props.title = f"{name} - Resume"
        core_props.subject = "Resume / CV"

        # Combine skills for keywords (limited to 255 chars for DOCX standard)
        keywords = []
        for cat, items in skills_dict.items():
            if items:
                # Only add category name to stay within limit
                keywords.extend(items.split(',')[:3])  # Take first 3 items per category

        keyword_str = ", ".join(keywords)[:250]  # Limit to 250 chars to be safe
        core_props.keywords = keyword_str
        
        # Use first line of summary as comments/description if available
        if summary:
            core_props.comments = summary[:200]

    def add_personal_info(self, name, email, phone, city, country, linkedin, github, summary):
        """Add name, contact info, and professional summary"""
        # Name - Heading 0 with optimized styling (20pt, dark blue)
        h = self.doc.add_heading(name, level=0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for run in h.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(20)  # Professional size
            run.font.color.rgb = RGBColor(25, 55, 95)  # Dark blue

        # Add spacing after name
        self.doc.add_paragraph()

        # Contact info with bullet separators
        contact_parts = []
        if email:
            contact_parts.append(email)
        if phone:
            contact_parts.append(phone)
        if city and country:
            if self.language == 'en':
                country = country.replace('Türkiye', 'Turkey')
            contact_parts.append(f"{city}, {country}")

        if contact_parts:
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("  •  ".join(contact_parts))
            run.font.size = Pt(10)

        # Links with subtle gray color
        if linkedin or github:
            links = []
            if linkedin:
                links.append(linkedin)
            if github:
                links.append(github)
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("  •  ".join(links))
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(80, 80, 80)  # Gray for links

        # Summary section
        if summary:
            self.doc.add_paragraph()  # Spacing
            self.add_section_title(self.labels[self.language]['SUMMARY'])
            p = self.doc.add_paragraph(summary)
            # Improved line spacing
            p.paragraph_format.line_spacing = 1.15
            self.doc.add_paragraph()  # Spacing

    def add_section_title(self, title):
        """Add section header using Heading 1 for semantic structure with modern styling"""
        # Add spacing before section
        self.doc.add_paragraph()

        h = self.doc.add_heading(title.upper(), level=1)

        for run in h.runs:
            run.font.name = 'Calibri'
            run.font.size = Pt(13)  # Slightly larger for better hierarchy
            run.font.bold = True
            run.font.color.rgb = RGBColor(25, 55, 95)  # Professional dark blue

        # Add subtle bottom border effect (using a paragraph with border)
        p = self.doc.add_paragraph()
        p_format = p.paragraph_format
        p_format.space_after = Pt(8)
        p_format.space_before = Pt(0)

    def add_experience_item(self, title, company, location, start_date, end_date, description):
        """Add a job/experience entry with improved styling"""
        # Job title with optimized formatting
        p = self.doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

        # Company and dates on same line
        # Company, Location and dates
        p = self.doc.add_paragraph()
        
        # Company & Location
        comp_loc = company
        if location:
            if self.language == 'en':
                location = location.replace('Türkiye', 'Turkey')
            comp_loc += f", {location}"
            
        run = p.add_run(comp_loc)
        run.font.size = Pt(10)
        run.italic = True
        
        # Translate Date
        if self.language == 'en' and end_date == 'Devam Ediyor':
            end_date = 'Present'
            
        # Dates with subtle gray color
        date_str = f"  |  {start_date} - {end_date}"
        run2 = p.add_run(date_str)
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(100, 100, 100)  # Gray for dates

        # Description with semantic list bullets and improved spacing
        if description:
            lines = description.strip().split('\n')
            for line in lines:
                line = line.strip()
                if line:
                    try:
                        p = self.doc.add_paragraph(line, style='List Bullet')
                        p.paragraph_format.line_spacing = 1.15
                    except KeyError:
                        p = self.doc.add_paragraph(f"• {line}")
                        p.paragraph_format.line_spacing = 1.15

        self.doc.add_paragraph()  # Spacing after entry

    def add_education_item(self, degree, school, year):
        """Add education entry with improved styling"""
        # Degree - Bold
        p = self.doc.add_paragraph()
        run = p.add_run(degree)
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Calibri'

        # School and year on same line
        p = self.doc.add_paragraph()
        run = p.add_run(school)
        run.font.size = Pt(10)

        # Year with subtle gray color
        run2 = p.add_run(f"  |  {year}")
        run2.font.size = Pt(9)
        run2.font.color.rgb = RGBColor(100, 100, 100)

        self.doc.add_paragraph()  # Spacing

    def add_skills(self, skills_dict):
        """Add skills section with improved formatting"""
        self.add_section_title(self.labels[self.language]['SKILLS'])

        # Translation Map
        cat_map = {
            'Programlama': 'Programming',
            'Frameworks': 'Frameworks',
            'Araçlar': 'Tools',
            'Diller': 'Languages'
        } if self.language == 'en' else {}

        for category, items in skills_dict.items():
            category = cat_map.get(category, category)
            if items:
                p = self.doc.add_paragraph()
                # Category in dark blue
                run = p.add_run(f"{category}: ")
                run.bold = True
                run.font.color.rgb = RGBColor(25, 55, 95)

                # Items in black
                run2 = p.add_run(items)
                run2.font.size = Pt(10)

                # Improved spacing between skill categories
                p.paragraph_format.space_after = Pt(4)

    def generate(self, data):
        """Generate complete resume from data dict"""
        # Personal info
        p = data.get('personal', {})
        
        # Set Metadata first
        self.set_metadata(
            p.get('fullName', 'Resume'), 
            p.get('summary', ''), 
            data.get('skills', {})
        )
        
        self.add_personal_info(
            p.get('fullName', ''),
            p.get('email', ''),
            p.get('phone', ''),
            p.get('city', ''),
            p.get('country', ''),
            p.get('linkedin', ''),
            p.get('github', ''),
            p.get('summary', '')
        )

        # Experience
        if data.get('experience'):
            self.add_section_title(self.labels[self.language]['EXPERIENCE'])
            for exp in data['experience']:
                self.add_experience_item(
                    exp.get('title', ''),
                    exp.get('company', ''),
                    exp.get('location', ''),
                    exp.get('startDate', ''),
                    exp.get('endDate', ''),
                    exp.get('description', '')
                )

        # Education
        if data.get('education'):
            self.add_section_title(self.labels[self.language]['EDUCATION'])
            for edu in data['education']:
                self.add_education_item(
                    edu.get('degree', ''),
                    edu.get('school', ''),
                    edu.get('year', '')
                )

        # Skills
        if data.get('skills'):
            self.add_skills(data['skills'])

    def save(self, filename):
        """Save document to file"""
        self.doc.save(filename)

    def get_bytes(self):
        """Return document as bytes for streaming"""
        buffer = BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()
